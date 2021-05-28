import docx
from os.path import join

origin_name = "data/origin_data.docx"
no_empty_name = "data/no_empty_line_data.docx"

#去除空行
def enter_remove():
    #获取文档对象
    in_file = docx.Document(origin_name)
    out_file = docx.Document()

    # 逐段处理
    for para in in_file.paragraphs:
        if len(para.text) == 0:
            continue
        out_file.add_paragraph(para.text)

    out_file.save(no_empty_name)

#找到各病例检测阳性日期，返回一个字典，key为日期，内容为当天检测阳性的病例号
def get_positive_data():

    file = docx.Document(no_empty_name)
    date_num = {}#结果字典
    num = 0
    date = ''

    for para in file.paragraphs:
        if len(para.text) < 20:#病例头，可读出序号
            for i in para.text:
                if '0' <= i <= '9':
                    num = num * 10 + int(i)
            #print(num)

        else:
            for i in range(len(para.text)):
                if para.text[i] == '阳' and para.text[i + 1] == '性':#探测到“阳性”
                    #print('find')
                    for back in range(i,0,-1):
                        if para.text[back] == '日' and '0' <= para.text[back - 1] <= '9':#往前搜，搜到数字+“日”
                            for getdate in range(back,0,-1):
                                if para.text[getdate] ==  '，' or para.text[getdate] =='。' or para.text[getdate] == '；' or para.text[getdate] ==  ',' or para.text[getdate] ==  ';':#搜到标点结束
                                    date_num.setdefault(date,[]).append(num)
                                    # print(num)
                                    # print(date)
                                    date = ''
                                    break
                                date = para.text[getdate] + date
                            break
                    break
            num = 0

    #with open("data/date_num.txt","w") as out_file:
    for i in date_num.items():
        print(i)
        #out_file.write(i)
        #out_file.write('\r\n')


#找到各病例居住地，得到一个字典，key为病例号，内容为住址
def get_address():
    file = docx.Document(no_empty_name)
    num_address = {}#结果字典
    num = 0
    comma = 0
    address = ""

    for para in file.paragraphs:
        if len(para.text) < 20:#病例头，可读出序号
            for i in para.text:
                if '0' <= i <= '9':
                    num = num * 10 + int(i)
            #print(num)
        else:
            for i in range(len(para.text)):
                if para.text[i] == '，' or para.text[i] == ',':
                    comma = comma + 1
                if comma == 2:
                    for getAddress in range(i + 1,len(para.text),1):

                        if para.text[getAddress] == "。" or para.text[getAddress] == "人" or para.text[getAddress] == "，":
                            num_address[num] = address
                            break
                        address = address + para.text[getAddress]
                    #print(address)
                    address = ""
                    break
            num = 0
            comma = 0

    for i in num_address.items():
        print(i)

if __name__ == '__main__':
   #enter_remove()
    #get_positive_data()
   get_address()
